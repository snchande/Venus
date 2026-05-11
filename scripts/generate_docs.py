"""
Venus Notebooks - Documentation Generator
Creates a professional Word document (.docx) for the Venus Notebooks project.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Colour palette ──────────────────────────────────────────────
VENUS_PURPLE   = RGBColor(0x8B, 0x5C, 0xF6)   # brand accent
VENUS_BLUE     = RGBColor(0x1D, 0x4E, 0xD8)   # headings
DARK_NAVY      = RGBColor(0x1E, 0x1E, 0x3F)   # title
MID_GRAY       = RGBColor(0x4B, 0x55, 0x63)   # body text
LIGHT_GRAY     = RGBColor(0xF3, 0xF4, 0xF6)   # table fills
DARK_PURPLE_BG = RGBColor(0x3B, 0x0F, 0x6E)   # cover background strip
WHITE          = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG        = RGBColor(0xF0, 0xF0, 0xF0)

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "..", "docs",
                           "Venus_Notebooks_Documentation.docx")


# ── Helpers ──────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="C7D2FE"):
    """Set thin coloured border on all sides of a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_horizontal_rule(doc, color="8B5CF6", thickness=6):
    """Insert a coloured horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading(doc, text, level=1, color=VENUS_BLUE, space_before=18, space_after=6):
    """Add a styled heading."""
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after  = Pt(space_after)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
    return h


def body(doc, text, color=MID_GRAY, italic=False, size=11):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.font.italic = italic
    return p


def bullet(doc, text, color=MID_GRAY, bold_prefix=None):
    """Add a bullet point; optional bold_prefix before colon."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.color.rgb = VENUS_BLUE
        remaining = text
        r2 = p.add_run(remaining)
        r2.font.color.rgb = color
        r2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.color.rgb = color
        run.font.size = Pt(11)
    return p


def code_block(doc, code_text):
    """Render a grey monospace code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EFEFEF")
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p


def two_col_table(doc, rows_data, header=None, col_widths=(2.2, 4.3)):
    """Create a two-column styled table."""
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    if header:
        row = table.add_row()
        for i, h_text in enumerate(header):
            cell = row.cells[i]
            cell.width = Inches(col_widths[i])
            set_cell_bg(cell, "3B0F6E")
            p = cell.paragraphs[0]
            run = p.add_run(h_text)
            run.font.bold  = True
            run.font.color.rgb = WHITE
            run.font.size  = Pt(10)

    for idx, (col1, col2) in enumerate(rows_data):
        row = table.add_row()
        bg = "F5F3FF" if idx % 2 == 0 else "FFFFFF"
        for i, txt in enumerate((col1, col2)):
            cell = row.cells[i]
            cell.width = Inches(col_widths[i])
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(10)
            run.font.color.rgb = MID_GRAY
            if i == 0:
                run.font.bold = True
                run.font.color.rgb = VENUS_BLUE

    doc.add_paragraph()
    return table


def page_break(doc):
    doc.add_page_break()


# ── Build Document ────────────────────────────────────────────────

def build():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════

    # Top accent bar (simulated with a coloured paragraph)
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "3B0F6E")
        pPr.append(shd)
        run = p.add_run(" " * 120)
        run.font.size = Pt(6)

    doc.add_paragraph()

    # Logo / title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(32)
    title_p.paragraph_format.space_after  = Pt(4)
    r = title_p.add_run("Venus Notebooks")
    r.font.size = Pt(36)
    r.font.bold = True
    r.font.color.rgb = DARK_NAVY

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(2)
    r2 = subtitle_p.add_run("Interactive Java Notebook Environment")
    r2.font.size = Pt(16)
    r2.font.color.rgb = VENUS_PURPLE
    r2.font.italic = True

    add_horizontal_rule(doc)

    doc.add_paragraph()

    meta_lines = [
        ("Project",        "Venus Notebooks v1.0"),
        ("Created by",     "Suresh Chande"),
        ("First Release",  "March 8, 2026"),
        ("Document Date",  "March 8, 2026"),
        ("Document Type",  "Technical Documentation"),
        ("Status",         "Initial Release"),
    ]
    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in meta_lines:
        row = meta_table.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(1.8)
        c1.width = Inches(3.6)
        set_cell_bg(c0, "EDE9FE")
        set_cell_borders(c0)
        set_cell_borders(c1)
        r0 = c0.paragraphs[0].add_run(label)
        r0.font.bold = True
        r0.font.color.rgb = VENUS_BLUE
        r0.font.size = Pt(11)
        r1 = c1.paragraphs[0].add_run(value)
        r1.font.color.rgb = MID_GRAY
        r1.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    tagline_p = doc.add_paragraph()
    tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = tagline_p.add_run(
        "Powered by JShell · Spring Boot · Claude AI · Maven Central"
    )
    rt.font.size = Pt(10)
    rt.font.color.rgb = RGBColor(0xA0, 0xA0, 0xC0)
    rt.font.italic = True

    # Bottom accent bar
    doc.add_paragraph()
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "8B5CF6")
        pPr.append(shd)
        run = p.add_run(" " * 120)
        run.font.size = Pt(6)

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ════════════════════════════════════════════════════════════
    heading(doc, "Table of Contents", level=1, space_before=0)
    add_horizontal_rule(doc)
    toc_items = [
        ("1", "Project Overview",                     "3"),
        ("2", "Creator & Version History",            "4"),
        ("3", "Technology Stack",                     "4"),
        ("4", "Features",                             "5"),
        ("5", "Architecture",                         "6"),
        ("6", "Project Structure",                    "7"),
        ("7", "Getting Started",                      "8"),
        ("8", "User Interface Guide",                 "10"),
        ("9", "REST API Reference",                   "12"),
        ("10","Package Manager",                      "15"),
        ("11","AI Assistant (Claude Integration)",    "16"),
        ("12","Notebook File Format",                 "17"),
        ("13","Configuration Reference",              "18"),
        ("14","Security Considerations",              "19"),
        ("15","Troubleshooting",                      "19"),
        ("16","Roadmap & Future Enhancements",        "20"),
        ("17","License",                              "21"),
    ]
    toc_table = doc.add_table(rows=0, cols=3)
    toc_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for num, title, pg in toc_items:
        row = toc_table.add_row()
        for i, (txt, w) in enumerate(zip((num, title, pg), (0.4, 5.0, 0.7))):
            cell = row.cells[i]
            cell.width = Inches(w)
            p2 = cell.paragraphs[0]
            run = p2.add_run(txt)
            run.font.size = Pt(11)
            if i == 0:
                run.font.bold = True
                run.font.color.rgb = VENUS_PURPLE
            elif i == 2:
                run.font.color.rgb = MID_GRAY
                p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                run.font.color.rgb = MID_GRAY

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 1. PROJECT OVERVIEW
    # ════════════════════════════════════════════════════════════
    heading(doc, "1. Project Overview", level=1)
    add_horizontal_rule(doc)

    body(doc,
        "Venus Notebooks is a locally-hosted, browser-based interactive Java development environment "
        "inspired by Jupyter Notebooks. It allows developers to write, execute, and document Java code "
        "in an interactive, cell-based format — without any cloud dependencies or complex setup.")

    body(doc,
        "The application is powered by JShell (Java's official REPL introduced in Java 9), which enables "
        "live, incremental execution of Java code. Results appear instantly in the browser through a "
        "WebSocket-driven real-time pipeline. A built-in Maven package manager lets users install any "
        "library from Maven Central on demand, while an integrated Claude AI assistant provides intelligent "
        "code generation, explanation, and notebook creation.")

    heading(doc, "Key Highlights", level=2)
    highlights = [
        ("Browser-based Java REPL", " — Execute Java code interactively in your browser, powered by JShell"),
        ("Real-time Output",        " — Console output streams instantly via STOMP WebSocket"),
        ("Maven Package Manager",   " — Install any Maven Central package without leaving the app"),
        ("AI-powered",              " — Integrated Claude API for code help and notebook generation"),
        ("Zero build step UI",      " — Pure HTML/CSS/JavaScript frontend, no webpack or npm"),
        ("Local-first",             " — All data stays on your machine; no cloud account required"),
        ("GitHub-ready",            " — Clean project structure with full documentation"),
    ]
    for bold, rest in highlights:
        p = doc.add_paragraph(style="List Bullet")
        r1 = p.add_run(bold)
        r1.font.bold = True
        r1.font.color.rgb = VENUS_BLUE
        r1.font.size = Pt(11)
        r2 = p.add_run(rest)
        r2.font.color.rgb = MID_GRAY
        r2.font.size = Pt(11)

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 2. CREATOR & VERSION HISTORY
    # ════════════════════════════════════════════════════════════
    heading(doc, "2. Creator & Version History", level=1)
    add_horizontal_rule(doc)

    heading(doc, "Creator", level=2)
    two_col_table(doc, [
        ("Name",          "Suresh Chande"),
        ("Role",          "Creator & Lead Developer"),
        ("Project",       "Venus Notebooks"),
        ("First Release", "March 8, 2026"),
        ("Contact",       "GitHub: github.com/yourusername"),
    ], header=("Field", "Details"))

    heading(doc, "Version History", level=2)
    two_col_table(doc, [
        ("v1.0.0  —  March 8, 2026",
         "Initial release. JShell execution, WebSocket streaming, Maven package manager, "
         "Claude AI integration, Settings panel, dark/light themes, .vnb notebook format."),
    ], header=("Version", "Changes"), col_widths=(2.4, 4.1))

    # ════════════════════════════════════════════════════════════
    # 3. TECHNOLOGY STACK
    # ════════════════════════════════════════════════════════════
    heading(doc, "3. Technology Stack", level=1)
    add_horizontal_rule(doc)

    two_col_table(doc, [
        ("Java 21 (JDK)",        "Runtime language and platform. Java 21 LTS is recommended."),
        ("Spring Boot 3.2",      "Embedded web server (Tomcat), REST MVC, WebSocket (STOMP)."),
        ("JShell API",           "jdk.jshell module — interactive Java code execution engine."),
        ("STOMP / SockJS",       "WebSocket protocol for real-time browser ↔ server communication."),
        ("Maven",                "Build tool and dependency management (pom.xml)."),
        ("Maven Central HTTP",   "Packages downloaded via HTTP from repo1.maven.org at runtime."),
        ("Anthropic Claude API", "REST API for AI chat, code generation, and notebook creation."),
        ("CodeMirror 5",         "Browser-based code editor with Java syntax highlighting (CDN)."),
        ("Marked.js",            "Markdown rendering for documentation cells (CDN)."),
        ("Lombok",               "Boilerplate reduction (@Data, @Builder) for Java model classes."),
        ("Jackson",              "JSON serialisation/deserialisation for notebooks and settings."),
        ("WebFlux WebClient",    "Reactive HTTP client used for Claude API calls."),
    ], header=("Technology", "Purpose"))

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 4. FEATURES
    # ════════════════════════════════════════════════════════════
    heading(doc, "4. Features", level=1)
    add_horizontal_rule(doc)

    features = [
        ("Notebook Editor",
         ["Code cells — execute Java code using JShell",
          "Markdown cells — write rich documentation with headers, lists, code blocks",
          "Cell operations: run, move up/down, toggle type, delete",
          "Run All Cells in sequence",
          "Per-notebook JShell sessions (variables isolated between notebooks)",
          "Auto-save support (configurable interval)",
          "Ctrl+S keyboard shortcut for saving"]),
        ("Interactive Console",
         ["Full JShell REPL in the browser",
          "Command history navigation (↑/↓ arrows)",
          "Shift+Enter for multi-line input",
          "Separate session from notebook (session ID configurable)",
          "Clear and restart controls"]),
        ("Package Manager",
         ["Install Maven packages by groupId:artifactId:version",
          "Download JARs directly from Maven Central",
          "Immediate classpath injection into all active JShell sessions",
          "Search Maven Central without leaving the app",
          "Persistent package list across restarts (data/packages.json)",
          "Remove installed packages"]),
        ("Claude AI Integration",
         ["Chat interface with conversation history",
          "Quick action buttons for common queries",
          "Generate complete notebooks from natural language prompts",
          "Explain Java code snippets",
          "Suggest fixes for JShell errors",
          "Configurable model (Sonnet, Opus, Haiku) and token limit"]),
        ("Settings & Admin",
         ["Claude API key management (masked display, env variable support)",
          "Theme switching: dark (default) / light",
          "Editor font size and line numbers toggle",
          "Auto-save interval configuration",
          "Live server status panel (Java version, API key status)"]),
    ]

    for feat_name, feat_bullets in features:
        heading(doc, feat_name, level=2)
        for b in feat_bullets:
            bullet(doc, b)

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 5. ARCHITECTURE
    # ════════════════════════════════════════════════════════════
    heading(doc, "5. Architecture", level=1)
    add_horizontal_rule(doc)

    body(doc,
        "Venus Notebooks follows a classic single-server architecture. One Spring Boot process "
        "serves both the static web UI and the REST/WebSocket API. There is no separate frontend "
        "build server, no database, and no cloud dependency.")

    heading(doc, "Request Flow", level=2)
    body(doc, "1.  Browser loads index.html + CSS + JS from Spring's static resource handler.")
    body(doc, "2.  JavaScript establishes a STOMP-over-SockJS WebSocket connection to /ws.")
    body(doc, "3.  User actions trigger REST calls to /api/* endpoints.")
    body(doc, "4.  Shell execution results are broadcast over WebSocket to /topic/shell/{sessionId}.")
    body(doc, "5.  Claude API calls are proxied server-side (API key never exposed to browser).")

    heading(doc, "JShell Session Management", level=2)
    body(doc,
        "Each notebook gets its own ShellSession — a wrapper around a jdk.jshell.JShell instance. "
        "Sessions live in memory for the server process lifetime. Variables accumulate within a session "
        "across cell executions, but are isolated between notebooks. The interactive Console tab uses "
        "a dedicated 'console' session.")

    heading(doc, "Data Storage", level=2)
    two_col_table(doc, [
        ("notebooks/*.vnb",       "Notebook files in JSON format. Each file is one notebook."),
        ("data/settings.json",    "Application settings including Claude API key (gitignored)."),
        ("data/packages.json",    "List of installed Maven packages with metadata."),
        ("data/packages/*.jar",   "Downloaded JAR files for installed packages."),
    ], header=("Location", "Contents"))

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 6. PROJECT STRUCTURE
    # ════════════════════════════════════════════════════════════
    heading(doc, "6. Project Structure", level=1)
    add_horizontal_rule(doc)

    code_block(doc, """\
venus/
├── CLAUDE.md                        # Claude Code AI instructions
├── README.md                        # User-facing documentation
├── pom.xml                          # Maven build file (Spring Boot 3.2.3)
├── .gitignore
│
├── .claude/
│   ├── settings.json                # Claude Code IDE settings
│   └── commands/
│       ├── start.md                 # /start slash command
│       └── build.md                 # /build slash command
│
├── docs/
│   ├── ARCHITECTURE.md
│   ├── API.md                       # Full REST API reference
│   ├── SETUP.md                     # Installation guide
│   ├── USAGE.md                     # End-user guide
│   └── Venus_Notebooks_Documentation.docx
│
├── src/main/java/com/venus/
│   ├── VenusApplication.java        # Spring Boot entry point
│   ├── config/
│   │   ├── WebSocketConfig.java     # STOMP WebSocket setup
│   │   └── CorsConfig.java
│   ├── model/                       # Data classes (Lombok @Data)
│   │   ├── Notebook.java
│   │   ├── Cell.java
│   │   ├── CellType.java
│   │   ├── PackageInfo.java
│   │   ├── VenusSettings.java
│   │   └── ExecutionResult.java
│   ├── shell/                       # JShell engine
│   │   ├── JShellManager.java       # Session registry + WebSocket broadcast
│   │   └── ShellSession.java        # Single JShell instance wrapper
│   ├── service/                     # Business logic
│   │   ├── NotebookService.java
│   │   ├── PackageService.java
│   │   ├── ClaudeService.java
│   │   └── SettingsService.java
│   └── controller/                  # REST endpoints
│       ├── NotebookController.java  # /api/notebooks
│       ├── ShellController.java     # /api/shell + WebSocket @MessageMapping
│       ├── PackageController.java   # /api/packages
│       ├── LLMController.java       # /api/llm
│       └── SettingsController.java  # /api/settings
│
├── src/main/resources/
│   ├── application.properties       # Port 8585, directory paths
│   └── static/
│       ├── index.html               # Single-page application shell
│       ├── css/venus.css            # Dark + light themes, all styles
│       └── js/
│           ├── app.js               # Tab navigation, WebSocket, REST helpers
│           ├── notebook.js          # Notebook editor + cell execution
│           ├── console-tab.js       # Interactive REPL console
│           ├── packages.js          # Package manager UI
│           ├── settings.js          # Settings form
│           └── ai-assistant.js      # Claude chat + generation
│
├── notebooks/
│   └── welcome.vnb                  # Sample notebook (checked in to git)
├── data/
│   ├── packages.json                # Installed packages list
│   └── packages/                   # Downloaded JARs
└── scripts/
    ├── start.sh                     # Unix/Mac startup script
    └── start.bat                    # Windows startup script""")

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 7. GETTING STARTED
    # ════════════════════════════════════════════════════════════
    heading(doc, "7. Getting Started", level=1)
    add_horizontal_rule(doc)

    heading(doc, "Prerequisites", level=2)
    two_col_table(doc, [
        ("Java JDK 17+",   "Full JDK required (not JRE). JDK 21 LTS is recommended."),
        ("Apache Maven",   "Version 3.8 or higher."),
        ("Internet",       "For Claude API calls and Maven Central package downloads."),
        ("Modern Browser", "Chrome, Firefox, Edge, or Safari (WebSocket required)."),
    ], header=("Requirement", "Notes"))

    heading(doc, "Quick Start", level=2)
    body(doc, "Step 1 — Clone or download the project:")
    code_block(doc, "git clone https://github.com/yourusername/venus-notebooks.git\ncd venus-notebooks")

    body(doc, "Step 2 — Build the project:")
    code_block(doc, "mvn clean package -DskipTests")

    body(doc, "Step 3 — Start Venus Notebooks:")
    code_block(doc, "# Windows\nscripts\\start.bat\n\n# Unix / macOS\n./scripts/start.sh")

    body(doc, "Step 4 — Open your browser:")
    code_block(doc, "http://localhost:8585")

    heading(doc, "Required JVM Flags", level=2)
    body(doc,
        "JShell is in a restricted JDK module (jdk.jshell). The following flags must always "
        "be passed to the JVM when running Venus Notebooks:")
    code_block(doc, """\
--add-opens=jdk.jshell/jdk.jshell=ALL-UNNAMED
--add-opens=java.base/java.lang=ALL-UNNAMED
--add-exports=jdk.jshell/jdk.jshell=ALL-UNNAMED""")
    body(doc, "The startup scripts and Maven plugin configuration already include these flags automatically.")

    heading(doc, "Configuring the Claude API Key", level=2)
    body(doc,
        "AI features require an Anthropic API key. You can configure it in three ways:")
    bullet(doc, "Settings tab in the UI: paste your key, click Save Settings")
    bullet(doc, "Environment variable (recommended for security):")
    code_block(doc, "# Windows\nset ANTHROPIC_API_KEY=sk-ant-your-key-here\n\n# Unix/Mac\nexport ANTHROPIC_API_KEY=sk-ant-your-key-here")
    bullet(doc, "Directly in data/settings.json (do NOT commit this file to git)")

    body(doc, "Get an API key at: https://console.anthropic.com/", italic=True)

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 8. USER INTERFACE GUIDE
    # ════════════════════════════════════════════════════════════
    heading(doc, "8. User Interface Guide", level=1)
    add_horizontal_rule(doc)

    heading(doc, "Tab Overview", level=2)
    two_col_table(doc, [
        ("Notebook",   "Main editor with code and markdown cells. Run Java interactively."),
        ("Console",    "Interactive JShell REPL — type and execute Java statements directly."),
        ("Packages",   "Install/remove Maven packages. Search Maven Central."),
        ("Settings",   "Configure Claude API key, theme, font size, auto-save, and more."),
        ("AI",         "Claude AI chat, quick actions, and AI-powered notebook generation."),
    ], header=("Tab", "Purpose"))

    heading(doc, "Notebook Tab", level=2)
    body(doc,
        "The Notebook tab is the main workspace. A notebook contains an ordered list of cells. "
        "Cells come in two types:")
    bullet(doc, "CODE — contains executable Java code run by JShell")
    bullet(doc, "MARKDOWN — contains documentation text rendered as formatted HTML")

    body(doc, "Cell operations:")
    two_col_table(doc, [
        ("Run (Shift+Enter / Ctrl+Enter)", "Execute the cell in the current JShell session"),
        ("+ Code / + Markdown",           "Add a new cell at the bottom"),
        ("↑ / ↓ arrows",                  "Move a cell up or down"),
        ("⇄ (toggle type)",               "Switch between CODE and MARKDOWN"),
        ("✕ (delete)",                    "Remove the cell"),
        ("Run All",                        "Execute all CODE cells sequentially"),
        ("↺ Restart",                     "Clear all JShell variables (re-run cells to restore state)"),
        ("Clear",                          "Clear all cell outputs"),
        ("Save (Ctrl+S)",                  "Save the notebook to disk"),
    ], header=("Action", "Description"), col_widths=(2.4, 4.1))

    heading(doc, "Console Tab", level=2)
    body(doc,
        "The Console provides a standalone JShell REPL. Press Enter to run a statement, "
        "Shift+Enter for multi-line input. Use ↑/↓ to navigate command history. "
        "The Console uses its own JShell session, separate from notebook sessions.")

    heading(doc, "Keyboard Shortcuts", level=2)
    two_col_table(doc, [
        ("Shift+Enter / Ctrl+Enter", "Run the focused code cell"),
        ("Ctrl+S",                   "Save current notebook"),
        ("↑ / ↓ (Console)",         "Navigate command history"),
        ("Shift+Enter (Console)",    "Insert newline without executing"),
    ], header=("Shortcut", "Action"))

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 9. REST API REFERENCE
    # ════════════════════════════════════════════════════════════
    heading(doc, "9. REST API Reference", level=1)
    add_horizontal_rule(doc)

    body(doc, "Base URL: http://localhost:8585/api  —  All endpoints accept and return JSON.")

    heading(doc, "Notebooks API", level=2)
    two_col_table(doc, [
        ("GET    /api/notebooks",       "List all notebooks (metadata only)"),
        ("POST   /api/notebooks",       "Create a new notebook { name }"),
        ("GET    /api/notebooks/{id}",  "Get full notebook with all cells"),
        ("PUT    /api/notebooks/{id}",  "Save / update a notebook"),
        ("DELETE /api/notebooks/{id}",  "Delete a notebook"),
    ], header=("Endpoint", "Description"), col_widths=(2.8, 3.7))

    heading(doc, "Shell Execution API", level=2)
    two_col_table(doc, [
        ("POST   /api/shell/execute",           "Execute Java code — body: { sessionId, code, cellId }"),
        ("POST   /api/shell/{sessionId}/restart","Restart session (clear variables)"),
        ("DELETE /api/shell/{sessionId}",        "Close and remove a session"),
        ("GET    /api/shell/sessions",           "List active session IDs"),
        ("GET    /api/shell/{sessionId}/info",   "Session classpath and execution count"),
    ], header=("Endpoint", "Description"), col_widths=(2.8, 3.7))

    heading(doc, "Package Manager API", level=2)
    two_col_table(doc, [
        ("GET    /api/packages",                    "List installed packages"),
        ("POST   /api/packages/install",            "Install package { coordinate }"),
        ("DELETE /api/packages/{g}/{a}/{v}",        "Remove package by GAV"),
        ("GET    /api/packages/search?q=...",       "Search Maven Central"),
    ], header=("Endpoint", "Description"), col_widths=(2.8, 3.7))

    heading(doc, "AI / LLM API", level=2)
    two_col_table(doc, [
        ("POST /api/llm/chat",      "Chat with Claude { message, history }"),
        ("POST /api/llm/generate",  "Generate notebook { prompt }"),
        ("POST /api/llm/explain",   "Explain code { code }"),
        ("POST /api/llm/fix",       "Fix error { code, error }"),
    ], header=("Endpoint", "Description"), col_widths=(2.8, 3.7))

    heading(doc, "Settings API", level=2)
    two_col_table(doc, [
        ("GET /api/settings",        "Get settings (API key masked)"),
        ("PUT /api/settings",        "Update settings"),
        ("GET /api/settings/status", "Server status (Java version, API key set, etc.)"),
    ], header=("Endpoint", "Description"), col_widths=(2.8, 3.7))

    heading(doc, "WebSocket (STOMP)", level=2)
    body(doc, "Connect via SockJS at /ws, then use the STOMP protocol:")
    two_col_table(doc, [
        ("SUBSCRIBE /topic/shell/{sessionId}", "Receive ExecutionResult JSON after each execution"),
        ("SEND /app/shell/{sessionId}",        "Execute code { code, cellId }"),
    ], header=("Destination", "Purpose"), col_widths=(3.0, 3.5))

    heading(doc, "ExecutionResult Object", level=2)
    code_block(doc, """\
{
  "sessionId":       "nb-welcome",
  "cellId":          "cell-2",
  "output":          "Hello World\\n",
  "error":           "",
  "returnValue":     null,
  "status":          "VALID",
  "success":         true,
  "executionTimeMs": 45,
  "executionCount":  1
}""")

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 10. PACKAGE MANAGER
    # ════════════════════════════════════════════════════════════
    heading(doc, "10. Package Manager", level=1)
    add_horizontal_rule(doc)

    body(doc,
        "The Venus Package Manager enables installing any Maven Central library into your "
        "JShell sessions without leaving the browser. Packages are downloaded as JAR files "
        "and injected directly into the JShell classpath at runtime.")

    heading(doc, "Installing a Package", level=2)
    body(doc, "Navigate to the Packages tab and enter the Maven coordinate:")
    code_block(doc, "groupId:artifactId:version\n\nExamples:\n  com.google.code.gson:gson:2.10.1\n  org.apache.commons:commons-lang3:3.14.0\n  com.fasterxml.jackson.core:jackson-databind:2.16.1")
    body(doc, "Click Install. The JAR is downloaded from Maven Central and immediately available in all active JShell sessions.")

    heading(doc, "Using Installed Packages", level=2)
    code_block(doc, "// After installing gson:\nimport com.google.gson.Gson;\nvar gson = new Gson();\nvar json = gson.toJson(Map.of(\"project\", \"Venus\", \"version\", 1.0));\nSystem.out.println(json);")

    heading(doc, "How It Works", level=2)
    bullet(doc, "Download URL: https://repo1.maven.org/maven2/{group-path}/{artifactId}/{version}/...")
    bullet(doc, "JAR stored at: data/packages/{groupId}_{artifactId}_{version}.jar")
    bullet(doc, "Package list persisted: data/packages.json")
    bullet(doc, "On server restart, packages are re-applied to new JShell sessions automatically")

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 11. AI ASSISTANT
    # ════════════════════════════════════════════════════════════
    heading(doc, "11. AI Assistant (Claude Integration)", level=1)
    add_horizontal_rule(doc)

    body(doc,
        "Venus Notebooks integrates with Anthropic's Claude API to provide an intelligent "
        "Java coding assistant. All API calls are made server-side; the API key is never "
        "exposed to the browser.")

    heading(doc, "Chat Interface", level=2)
    body(doc,
        "The AI tab provides a conversational interface. Claude maintains context across "
        "messages in a session. Quick action buttons provide one-click access to common queries.")

    heading(doc, "Notebook Generation", level=2)
    body(doc,
        "Clicking Generate Notebook opens a prompt dialog. Enter a description of the notebook "
        "you want, and Claude will generate a complete .vnb notebook with markdown explanations "
        "and runnable Java code cells. The notebook is saved automatically and opened in the editor.")

    body(doc, "Example prompts:")
    bullet(doc, "\"A tutorial on Java streams with filtering, mapping, and collecting examples\"")
    bullet(doc, "\"Demonstrate all Java 21 features: records, sealed classes, pattern matching\"")
    bullet(doc, "\"Working with JSON data using Jackson ObjectMapper\"")

    heading(doc, "Claude API Configuration", level=2)
    two_col_table(doc, [
        ("Model",       "claude-sonnet-4-6 (default), opus-4-6, haiku-4-5"),
        ("Max Tokens",  "Default 4096, configurable up to 8192"),
        ("System Prompt","Pre-configured for Java + JShell context"),
        ("API Endpoint","https://api.anthropic.com/v1/messages"),
    ], header=("Setting", "Value / Notes"))

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 12. NOTEBOOK FILE FORMAT
    # ════════════════════════════════════════════════════════════
    heading(doc, "12. Notebook File Format", level=1)
    add_horizontal_rule(doc)

    body(doc,
        "Notebooks are stored as .vnb files (Venus Notebook) in JSON format. "
        "The format is designed to be human-readable, version-control-friendly, and "
        "easily shareable.")

    code_block(doc, """\
{
  "id":          "550e8400-e29b-41d4-a716-446655440000",
  "name":        "My Notebook",
  "description": "Optional description",
  "created":     "2026-03-08T00:00:00",
  "modified":    "2026-03-08T12:00:00",
  "cells": [
    {
      "id":             "cell-1",
      "type":           "MARKDOWN",
      "source":         "# Hello\\n\\nThis is a *markdown* cell.",
      "output":         "",
      "error":          "",
      "returnValue":    null,
      "executed":       false,
      "executionCount": null
    },
    {
      "id":             "cell-2",
      "type":           "CODE",
      "source":         "System.out.println(\\"Hello, Venus!\\");",
      "output":         "Hello, Venus!\\n",
      "error":          "",
      "returnValue":    null,
      "executed":       true,
      "executionCount": 1
    }
  ],
  "metadata": {
    "javaVersion": "21",
    "packages":    []
  }
}""")

    two_col_table(doc, [
        ("id",             "UUID or string identifier. Unique per notebook."),
        ("name",           "Display name shown in the notebook selector."),
        ("cells",          "Ordered array of Cell objects."),
        ("cell.type",      "\"CODE\" or \"MARKDOWN\""),
        ("cell.source",    "The raw source text (Java code or Markdown)."),
        ("cell.output",    "Captured stdout from last execution."),
        ("cell.error",     "Compile or runtime error text."),
        ("cell.returnValue","Last expression value (JShell return value)."),
        ("metadata",       "Free-form object for future extensions."),
    ], header=("Field", "Description"))

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 13. CONFIGURATION REFERENCE
    # ════════════════════════════════════════════════════════════
    heading(doc, "13. Configuration Reference", level=1)
    add_horizontal_rule(doc)

    heading(doc, "application.properties", level=2)
    two_col_table(doc, [
        ("server.port",                "8585 — HTTP server port"),
        ("venus.notebooks.dir",        "notebooks — Directory for .vnb files"),
        ("venus.data.dir",             "data — Directory for settings and packages"),
        ("spring.jackson.serialization\n.write-dates-as-timestamps",
                                       "false — ISO-8601 date format in JSON"),
    ], header=("Property", "Default / Description"))

    heading(doc, "VenusSettings (data/settings.json)", level=2)
    two_col_table(doc, [
        ("anthropicApiKey",     "Your Anthropic API key (sk-ant-...)"),
        ("claudeModel",         "claude-sonnet-4-6 (default)"),
        ("claudeMaxTokens",     "4096"),
        ("theme",               "dark or light"),
        ("editorFontSize",      "14 (pixels)"),
        ("showLineNumbers",     "true"),
        ("autoSaveIntervalSecs","30 (0 = disabled)"),
        ("maxExecutionTimeMs",  "30000 (30 seconds)"),
        ("maxOutputLines",      "1000"),
    ], header=("Setting", "Default"))

    heading(doc, "Environment Variables", level=2)
    two_col_table(doc, [
        ("ANTHROPIC_API_KEY", "Anthropic API key — takes priority over settings.json"),
        ("SERVER_PORT",       "Override server port"),
        ("VENUS_HOME",        "Project root directory (used by Claude Code)"),
    ], header=("Variable", "Purpose"))

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 14. SECURITY CONSIDERATIONS
    # ════════════════════════════════════════════════════════════
    heading(doc, "14. Security Considerations", level=1)
    add_horizontal_rule(doc)

    body(doc,
        "Venus Notebooks is designed for local, single-user use. The following points "
        "should be considered before exposing it to a network:")

    bullet(doc, "Code Execution — JShell runs with the same JVM permissions as the server process. "
                "Never expose Venus to untrusted users on a shared network.")
    bullet(doc, "API Keys — The Anthropic API key is stored in data/settings.json (gitignored). "
                "Use the ANTHROPIC_API_KEY environment variable for better security.")
    bullet(doc, "No Authentication — Venus has no login system. Use OS-level firewall rules "
                "to restrict access to localhost only.")
    bullet(doc, "CORS — Currently configured to allow all origins (suitable for local dev). "
                "Restrict CORS in CorsConfig.java if deploying to a server.")
    bullet(doc, "HTTPS — No TLS in the default setup. Run behind a reverse proxy (nginx) "
                "with TLS if network exposure is needed.")

    # ════════════════════════════════════════════════════════════
    # 15. TROUBLESHOOTING
    # ════════════════════════════════════════════════════════════
    heading(doc, "15. Troubleshooting", level=1)
    add_horizontal_rule(doc)

    two_col_table(doc, [
        ("InaccessibleObjectException on startup",
         "Missing JVM flags. Ensure all three --add-opens / --add-exports flags are present."),
        ("'javac' not found",
         "You have JRE, not JDK. Install JDK 17 or 21."),
        ("Port 8585 already in use",
         "Change server.port in application.properties or use --server.port=XXXX"),
        ("WebSocket red dot (disconnected)",
         "Check browser console. Ensure no proxy blocks WebSocket upgrades."),
        ("Package download fails HTTP 404",
         "Verify coordinates on search.maven.org. Version may not exist."),
        ("Claude API 401 Unauthorized",
         "API key is wrong or expired. Re-enter in Settings tab."),
        ("Claude API 529 Overloaded",
         "Anthropic API is busy. Wait a moment and retry."),
        ("Notebook won't save",
         "Check write permissions on the notebooks/ directory."),
    ], header=("Problem", "Solution"), col_widths=(2.6, 3.9))

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 16. ROADMAP
    # ════════════════════════════════════════════════════════════
    heading(doc, "16. Roadmap & Future Enhancements", level=1)
    add_horizontal_rule(doc)

    roadmap = [
        ("v1.1  —  Notebook Sharing",       [
            "Export notebooks as HTML or PDF",
            "Import Jupyter .ipynb notebooks",
            "Share notebooks via URL (base64-encoded or server link)"]),
        ("v1.2  —  Collaboration",          [
            "Multi-cursor collaborative editing (WebSocket-based)",
            "Basic user sessions / authentication",
            "Notebook commenting"]),
        ("v1.3  —  Enhanced Execution",     [
            "Execution timeout per cell",
            "Cell output truncation with expand button",
            "Inline chart / graph rendering (JavaFX snapshot to canvas)",
            "Dependency-aware cell execution order"]),
        ("v1.4  —  Deeper AI Integration",  [
            "Inline AI completions in the code editor",
            "Auto-fix errors directly in cells",
            "AI-generated test cells",
            "Context-aware package suggestions"]),
        ("v2.0  —  Plugin System",          [
            "Custom cell types (e.g. SQL, Groovy, Kotlin)",
            "Theme marketplace",
            "Extension API for third-party integrations"]),
    ]

    for version, items in roadmap:
        heading(doc, version, level=2)
        for item in items:
            bullet(doc, item)

    page_break(doc)

    # ════════════════════════════════════════════════════════════
    # 17. LICENSE
    # ════════════════════════════════════════════════════════════
    heading(doc, "17. License", level=1)
    add_horizontal_rule(doc)

    body(doc, "Venus Notebooks is released under the MIT License.")
    doc.add_paragraph()
    code_block(doc, """\
MIT License

Copyright (c) 2026 Suresh Chande

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in
all copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT.""")

    doc.add_paragraph()
    add_horizontal_rule(doc, color="8B5CF6")
    doc.add_paragraph()

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run("Venus Notebooks  ·  Created by Suresh Chande  ·  March 8, 2026")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xA0, 0xA0, 0xC0)
    r.font.italic = True

    # ── Save ──
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {os.path.abspath(OUTPUT_PATH)}")


if __name__ == "__main__":
    build()
